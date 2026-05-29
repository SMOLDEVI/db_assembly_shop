import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_page_number(run):
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

def build_docx():
    doc = docx.Document()

    # Настройка полей (ГОСТ: Левое 3см, Правое 1см, Верхнее 2см, Нижнее 2см)
    for section in doc.sections:
        section.top_margin = Inches(0.78)    # 2.0 cm
        section.bottom_margin = Inches(0.78) # 2.0 cm
        section.left_margin = Inches(1.18)   # 3.0 cm
        section.right_margin = Inches(0.39)  # 1.0 cm
        
        # Настройка нумерации страниц (различный колонтитул для первой страницы)
        section.different_first_page_header_footer = True
        
        # Добавляем нумерацию в нижний колонтитул
        footer = section.footer
        p_footer = footer.paragraphs[0]
        p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_footer.paragraph_format.first_line_indent = Inches(0)
        p_footer.paragraph_format.space_after = Pt(0)
        p_footer.paragraph_format.space_before = Pt(0)
        run_footer = p_footer.add_run()
        run_footer.font.name = 'Times New Roman'
        run_footer.font.size = Pt(11)
        add_page_number(run_footer)

    # Настройка дефолтного стиля Normal
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.first_line_indent = Inches(0.49) # Абзацный отступ 1.25 см

    # Вспомогательные функции форматирования
    def add_title_line(text, size=14, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0, space_before=0):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Inches(0)
        p.paragraph_format.alignment = align
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        return p

    def add_chapter(num_and_title):
        doc.add_page_break()
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Inches(0)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(12)
        run = p.add_run(num_and_title)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.bold = True
        return p

    def add_section(text):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Inches(0)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.bold = True
        return p

    # ================= 1. ТИТУЛЬНЫЙ ЛИСТ =================
    add_title_line("Федеральное государственное бюджетное образовательное учреждение", size=12)
    add_title_line("высшего образования", size=12)
    add_title_line("«РЫБИНСКИЙ ГОСУДАРСТВЕННЫЙ АВИАЦИОННЫЙ ТЕХНИЧЕСКИЙ УНИВЕРСИТЕТ", size=12, bold=True)
    add_title_line("ИМЕНИ П.А. СОЛОВЬЕВА»", size=12, bold=True, space_after=30)
    
    add_title_line("Факультет радиоэлектроники и информатики", size=12)
    add_title_line("Кафедра математического и программного обеспечения электронных вычислительных средств", size=11, space_after=100)

    add_title_line("КУРСОВОЙ ПРОЕКТ", size=20, bold=True, space_before=10)
    add_title_line("по дисциплине «Базы данных»", size=14, space_after=10)
    add_title_line("на тему: «Разработка базы данных и веб-приложения для автоматизации учета в сборочном цеху»", size=15, bold=True, space_after=150)
    
    add_title_line("ПОЯСНИТЕЛЬНАЯ ЗАПИСКА", size=14, bold=True, space_after=100)

    # Информация о руководителях и авторе
    p_info = doc.add_paragraph()
    p_info.paragraph_format.first_line_indent = Inches(0)
    p_info.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_info.paragraph_format.space_before = Pt(30)
    p_info.paragraph_format.space_after = Pt(80)
    p_info.paragraph_format.line_spacing = 1.15
    run_info = p_info.add_run(
        "Выполнил студент группы ИС-22\n"
        "Иванов И.И.\n\n"
        "Руководитель работы:\n"
        "Шаров В.Г., Задорина Н.А.\n"
    )
    run_info.font.name = 'Times New Roman'
    run_info.font.size = Pt(12)

    add_title_line("Рыбинск 2026", size=12, space_before=50)

    # ================= 2. ОГЛАВЛЕНИЕ =================
    doc.add_page_break()
    add_title_line("Оглавление", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
    
    # Оглавление в виде списка
    toc_items = [
        ("1. Введение", "3"),
        ("2. Постановка задачи", "5"),
        ("3. Описание предметной области", "6"),
        ("   3.1. Общее описание предметной области", "6"),
        ("   3.2. Описание входных документов", "7"),
        ("   3.3. Описание выходных документов", "7"),
        ("   3.4. Список ограничений", "7"),
        ("4. Инфологическая модель", "8"),
        ("   4.1. Сущности и их ключевые атрибуты", "8"),
        ("   4.2. Бизнес-правила", "8"),
        ("   4.3. ER-диаграмма", "9"),
        ("5. Даталогическая модель", "10"),
        ("   5.1. Отношения базы данных", "10"),
        ("   5.2. Описание структуры всех отношений БД", "10"),
        ("   5.3. Сводная таблица отношений", "11"),
        ("   5.4. Обоснование нормализации", "11"),
        ("6. Реализация базы данных", "12"),
        ("   6.1. Создание таблиц в БД", "12"),
        ("7. Реализация приложения", "13"),
        ("   7.1. Структура приложения", "13"),
        ("   7.2. Описание интерфейсных окон", "14"),
        ("8. Заключение", "22"),
        ("9. Список литературы", "23"),
        ("10. Приложения", "24")
    ]
    
    for item, page in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Inches(0)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        
        # Точки до номера страницы (выравнивание с помощью табуляции)
        dots_len = 80 - len(item) - len(page)
        dots = "." * dots_len if dots_len > 0 else " "
        
        run_item = p.add_run(f"{item} {dots} {page}")
        run_item.font.name = 'Times New Roman'
        run_item.font.size = Pt(12)

    # ================= 1. ВВЕДЕНИЕ =================
    add_chapter("1. Введение")
    doc.add_paragraph(
        "Актуальность разработки информационной системы «Сборочный цех» заключается в автоматизации "
        "ведения учета деталей, готовых изделий, спецификаций состава изделий, контроля за адресным складом "
        "как комплектующих, так и готовой продукции, а также создания инструментов планирования и выявления дефицита "
        "материалов с возможностью экспорта отчетов."
    )
    doc.add_paragraph(
        "Цель работы — спроектировать реляционную базу данных и разработать веб-приложение, которое упрощает "
        "контроль складских запасов и логику сборки сложных готовых изделий на основе спецификаций с автоматическим "
        "списанием деталей из ячеек хранения."
    )
    doc.add_paragraph(
        "Для достижения цели используются современные технологии веб-разработки: язык Python, микрофреймворк Flask "
        "для реализации логической части, СУБД SQLite для хранения данных и библиотека Chart.js для визуализации аналитики "
        "склада готовой продукции."
    )

    # ================= 2. ПОСТАНОВКА ЗАДАЧИ =================
    add_chapter("2. Постановка задачи")
    doc.add_paragraph(
        "Сборочный цех выпускает готовую продукцию из комплектующих деталей. Детали поступают на склад комплектующих, "
        "где размещаются в конкретных ячейках хранения. Готовая продукция собирается силами сборочных постов, "
        "при этом со склада списываются детали, входящие в состав изделия согласно спецификации, а собранный продукт "
        "помещается в ячейку склада готовой продукции."
    )
    doc.add_paragraph(
        "Информационная система должна обеспечивать выполнение следующих операций:\n"
        "• ведение справочников комплектующих деталей (CRUD);\n"
        "• ведение справочника готовых изделий и их спецификаций (составов);\n"
        "• учет адресного размещения деталей и готовой продукции на складах с указанием остатков;\n"
        "• автоматизация сборки изделий: проверка наличия деталей на складе, каскадное списание комплектующих из ячеек, зачисление готового изделия в ячейку;\n"
        "• расчет дефицита комплектующих под плановые показатели выпуска;\n"
        "• выгрузка отчетов об остатках и дефиците деталей в формате CSV."
    )

    # ================= 3. ОПИСАНИЕ ПРЕДМЕТНОЙ ОБЛАСТИ =================
    add_chapter("3. Описание предметной области")
    
    add_section("3.1. Общее описание предметной области")
    doc.add_paragraph(
        "В процессе функционирования предприятия начальник цеха или кладовщик взаимодействует со сборочным постом "
        "и складом. При поступлении деталей на склад они регистрируются в системе и раскладываются по ячейкам. "
        "Каждая деталь характеризуется названием и артикулом. Каждое изделие характеризуется названием и описанием."
    )
    doc.add_paragraph(
        "Связующим звеном является состав изделия. Для каждого изделия задается перечень входящих в него деталей с "
        "указанием количества каждой детали. При сборке изделия система проверяет суммарные остатки всех необходимых "
        "деталей по ячейкам. Если деталей хватает, система списывает их из ячеек (уменьшая остатки) и зачисляет "
        "готовое изделие в ячейку готовой продукции."
    )

    add_section("3.2. Описание входных документов")
    doc.add_paragraph(
        "Входными данными являются экранные формы добавления деталей, готовых изделий, создания спецификаций, "
        "а также формы регистрации ячеек складов деталей/изделий. Также входным документом служит форма указания "
        "плановых показателей для расчета дефицита."
    )

    add_section("3.3. Описание выходных документов")
    doc.add_paragraph(
        "Выходными документами служат табличные справки о текущих остатках на складах, инфографика распределения изделий "
        "в процентах на базе Chart.js, интерактивный отчет о дефиците деталей для выполнения плана и экспортируемые "
        "CSV-отчеты о состоянии склада и о дефиците материалов."
    )

    add_section("3.4. Список ограничений")
    doc.add_paragraph(
        "1. Размерность строковых полей (артикулы, наименования деталей и изделий, номера ячеек) — до 255 символов.\n"
        "2. Количество деталей в составе, остатки в ячейках и плановые показатели выпуска должны быть целыми положительными числами (от 0 до 2147483647).\n"
        "3. Одна ячейка склада деталей может содержать только одну деталь. Одна ячейка склада готовой продукции может содержать только одно изделие."
    )

    # ================= 4. ИНФОЛОГИЧЕСКАЯ МОДЕЛЬ =================
    add_chapter("4. Инфологическая модель")
    
    add_section("4.1. Сущности и их ключевые атрибуты")
    doc.add_paragraph(
        "В результате анализа предметной области были выделены следующие сущности:\n"
        "• Деталь (КодДетали, Название, Артикул);\n"
        "• Изделие (КодИзделия, Название, Описание);\n"
        "• Ячейка склада деталей (КодЯчейкиДеталей, НомерЯчейки, КодДетали, Количество);\n"
        "• Ячейка склада изделий (КодЯчейкиИзделий, НомерЯчейки, КодИзделия, Количество);\n"
        "• План производства (КодИзделия, ПланируемоеКоличество)."
    )

    add_section("4.2. Бизнес-правила")
    doc.add_paragraph(
        "• Изделие — Состав: Изделие может состоять из многих деталей. Деталь может входить в состав многих изделий.\n"
        "• Деталь — Ячейка склада: Деталь может быть размещена в нескольких ячейках. В ячейке может находиться только одна деталь.\n"
        "• Изделие — Ячейка готовой продукции: Изделие может находиться в нескольких ячейках готовой продукции. В ячейке хранится только одно изделие.\n"
        "• Изделие — План: Каждое изделие имеет один целевой показатель плана."
    )

    add_section("4.3. ER-диаграмма")
    doc.add_paragraph(
        "ER-диаграмма сущностей и связей представлена на рисунке 1."
    )
    add_title_line("[Рисунок 1 – ER-диаграмма связей сущностей базы данных «Сборочный цех»]", size=12, italic=True, space_before=12, space_after=12)

    # ================= 5. ДАТАЛОГИЧЕСКАЯ МОДЕЛЬ =================
    add_chapter("5. Даталогическая модель")
    
    add_section("5.1. Отношения базы данных")
    doc.add_paragraph(
        "Логическая схема базы данных представлена в виде следующих отношений (таблиц):\n"
        "• part(id_part, name, article);\n"
        "• product_type(id_product_type, name, description);\n"
        "• product_composition(product_type_id, part_id, quantity);\n"
        "• part_warehouse_cell(id_part_cell, cell_number, part_id, quantity_stored);\n"
        "• product_warehouse_cell(id_product_cell, cell_number, product_type_id, quantity_stored);\n"
        "• production_plan(product_type_id, target_quantity)."
    )

    add_section("5.2. Описание структуры всех отношений БД")
    
    # Таблица со структурой БД
    table_struct = doc.add_table(rows=1, cols=6)
    table_struct.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table_struct.rows[0].cells
    hdr[0].text = 'Инф. объект'
    hdr[1].text = 'Имя атрибута'
    hdr[2].text = 'Имя в БД'
    hdr[3].text = 'Тип'
    hdr[4].text = 'Длина'
    hdr[5].text = 'Ключ'
    
    for cell in hdr:
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
        cell.paragraphs[0].runs[0].font.size = Pt(11)

    fields_data = [
        ('Part', 'КодДетали', 'id_part', 'I', '20', '*'),
        ('Part', 'Название', 'name', 'S', '255', ''),
        ('Part', 'Артикул', 'article', 'S', '255', ''),
        ('Product', 'КодИзделия', 'id_product_type', 'I', '20', '*'),
        ('Product', 'Название', 'name', 'S', '255', ''),
        ('Product_Comp', 'КодИзделия', 'product_type_id', 'I', '20', '*'),
        ('Product_Comp', 'КодДетали', 'part_id', 'I', '20', '*'),
        ('Product_Comp', 'Количество', 'quantity', 'I', '20', ''),
        ('Part_Cell', 'КодЯчейки', 'id_part_cell', 'I', '20', '*'),
        ('Part_Cell', 'НомерЯчейки', 'cell_number', 'S', '255', ''),
        ('Part_Cell', 'Количество', 'quantity_stored', 'I', '20', ''),
        ('Prod_Plan', 'КодИзделия', 'product_type_id', 'I', '20', '*'),
        ('Prod_Plan', 'ПланВыпуска', 'target_quantity', 'I', '20', '')
    ]

    for item in fields_data:
        row = table_struct.add_row().cells
        for col_idx, text in enumerate(item):
            row[col_idx].text = text
            row[col_idx].paragraphs[0].runs[0].font.name = 'Times New Roman'
            row[col_idx].paragraphs[0].runs[0].font.size = Pt(10)

    doc.add_paragraph().paragraph_format.space_before = Pt(12)

    add_section("5.3. Сводная таблица отношений")
    
    table_summary = doc.add_table(rows=1, cols=2)
    table_summary.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_sum = table_summary.rows[0].cells
    hdr_sum[0].text = 'Имя таблицы'
    hdr_sum[1].text = 'Назначение таблицы в БД'
    for cell in hdr_sum:
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
        cell.paragraphs[0].runs[0].font.size = Pt(11)

    tables_desc = [
        ('part', 'Справочник комплектующих деталей с артикулами'),
        ('product_type', 'Справочник выпускаемых готовых изделий'),
        ('product_composition', 'Спецификации состава готовой продукции (расход деталей)'),
        ('part_warehouse_cell', 'Ячейки адресного хранения деталей на складе комплектующих'),
        ('product_warehouse_cell', 'Ячейки адресного хранения готовых изделий на складе продукции'),
        ('production_plan', 'Целевые показатели плана выпуска готовой продукции')
    ]
    
    for t_name, t_desc in tables_desc:
        row = table_summary.add_row().cells
        row[0].text = t_name
        row[1].text = t_desc
        row[0].paragraphs[0].runs[0].font.name = 'Times New Roman'
        row[0].paragraphs[0].runs[0].font.size = Pt(11)
        row[1].paragraphs[0].runs[0].font.name = 'Times New Roman'
        row[1].paragraphs[0].runs[0].font.size = Pt(11)

    doc.add_paragraph().paragraph_format.space_before = Pt(12)

    add_section("5.4. Обоснование нормализации")
    doc.add_paragraph(
        "Спроектированная база данных находится в третьей нормальной форме (3NF):\n"
        "• все атрибуты отношений являются атомарными (первая нормальная форма 1NF);\n"
        "• все неключевые атрибуты полностью зависят от первичного ключа и отсутствуют частичные функциональные зависимости от составных ключей (вторая нормальная форма 2NF);\n"
        "• отсутствуют транзитивные функциональные зависимости неключевых атрибутов от первичных ключей (третья нормальная форма 3NF)."
    )

    # ================= 6. РЕАЛИЗАЦИЯ БАЗЫ ДАННЫХ =================
    add_chapter("6. Реализация базы данных")
    doc.add_paragraph(
        "Разработка базы данных выполнена в реляционной СУБД SQLite. База данных содержит 6 таблиц. "
        "Связи контролируются внешними ключами с каскадным удалением."
    )
    
    add_section("6.1. Создание таблиц в БД")
    doc.add_paragraph(
        "Ниже приведены DDL-скрипты, использованные для генерации таблиц базы данных сборочного цеха:"
    )
    
    sql_script = (
        "CREATE TABLE part (\n"
        "    id_part INTEGER PRIMARY KEY AUTOINCREMENT,\n"
        "    name TEXT NOT NULL,\n"
        "    article TEXT NOT NULL\n"
        ");\n\n"
        "CREATE TABLE product_type (\n"
        "    id_product_type INTEGER PRIMARY KEY AUTOINCREMENT,\n"
        "    name TEXT NOT NULL,\n"
        "    description TEXT\n"
        ");\n\n"
        "CREATE TABLE product_composition (\n"
        "    product_type_id INTEGER,\n"
        "    part_id INTEGER,\n"
        "    quantity INTEGER,\n"
        "    PRIMARY KEY (product_type_id, part_id),\n"
        "    FOREIGN KEY (product_type_id) REFERENCES product_type (id_product_type) ON DELETE CASCADE,\n"
        "    FOREIGN KEY (part_id) REFERENCES part (id_part) ON DELETE CASCADE\n"
        ");\n\n"
        "CREATE TABLE part_warehouse_cell (\n"
        "    id_part_cell INTEGER PRIMARY KEY AUTOINCREMENT,\n"
        "    cell_number TEXT NOT NULL,\n"
        "    part_id INTEGER,\n"
        "    quantity_stored INTEGER DEFAULT 0,\n"
        "    FOREIGN KEY (part_id) REFERENCES part (id_part) ON DELETE CASCADE\n"
        ");"
    )
    
    p_code = doc.add_paragraph()
    p_code.paragraph_format.left_indent = Inches(0.25)
    p_code.paragraph_format.first_line_indent = Inches(0)
    p_code.paragraph_format.line_spacing = 1.0
    p_code.paragraph_format.space_before = Pt(6)
    p_code.paragraph_format.space_after = Pt(6)
    run_code = p_code.add_run(sql_script)
    run_code.font.name = 'Courier New'
    run_code.font.size = Pt(10)

    # ================= 7. РЕАЛИЗАЦИЯ ПРИЛОЖЕНИЯ =================
    add_chapter("7. Реализация приложения")
    
    add_section("7.1. Структура приложения")
    doc.add_paragraph(
        "Приложение построено по архитектуре MVC (Model-View-Controller) с использованием Flask:\n"
        "• app.py — контроллер приложения, содержащий описание маршрутов и бизнес-логику;\n"
        "• database.db — файл базы данных SQLite;\n"
        "• templates/ — каталог представлений (HTML-шаблоны Jinja2):\n"
        "  - base.html (базовый макет с навигацией);\n"
        "  - index.html (дашборд и инфографика);\n"
        "  - assemble.html (консоль автоматической сборки);\n"
        "  - reports.html (интерфейс планирования и дефицита);\n"
        "  - CRUD-шаблоны справочников (parts, products, compositions, cells)."
    )

    add_section("7.2. Описание интерфейсных окон")
    doc.add_paragraph(
        "1. Информационная панель (Дашборд). Выводит оперативные статистические карточки (количество изделий в наличии, "
        "количество деталей на складе, количество активных спецификаций, число дефицитных позиций) и диаграмму Chart.js "
        "для анализа распределения готовой продукции в процентном отношении.\n"
        "2. Справочники деталей и изделий. Представлены в виде таблиц со списками. Кнопка «+» открывает модальное окно "
        "Bootstrap с формой добавления элемента, валидация данных происходит на стороне сервера.\n"
        "3. Консоль автосборки. Дает возможность выбрать изделие, объем партии и указать ячейку. Алгоритм в app.py проверяет "
        "наличие деталей по цепочке спецификаций, списывает их из ячеек хранения и зачисляет готовое изделие.\n"
        "4. Страница дефицита и планирования. В левой части позволяет задать плановое количество выпуска изделий. В правой части "
        "выводит расчетную таблицу деталей, которых не хватает на складе для покрытия этого плана. Отсюда доступен экспорт в CSV."
    )

    # ================= 8. ЗАКЛЮЧЕНИЕ =================
    add_chapter("8. Заключение")
    doc.add_paragraph(
        "В ходе разработки курсового проекта были выполнены все этапы проектирования баз данных и программного обеспечения. "
        "Созданная СУБД сборочного цеха обеспечивает целостность данных за счет внешних ключей и ограничений. "
        "Алгоритм каскадного списания деталей из ячеек хранения автоматизирует сложный учет сборочных операций."
    )
    doc.add_paragraph(
        "Интерфейс планирования дефицита комплектующих позволяет оптимизировать снабжение цеха деталями. Приложение готово "
        "к эксплуатации и обеспечивает высокую скорость выполнения запросов."
    )

    # ================= 9. СПИСОК ЛИТЕРАТУРЫ =================
    add_chapter("9. Список литературы")
    
    lit_sources = [
        "Дейт К. Дж. Введение в системы баз данных. — 8-е изд. — М.: Издательский дом «Вильямс», 2006. — 1328 с.",
        "Григхем М. Flask на практике: разработка веб-приложений на Python. — М.: ДМК Пресс, 2020. — 280 с.",
        "Лутц М. Изучаем Python. Том 1. — 5-е изд. — СПб.: Диалектика, 2019. — 832 с.",
        "Официальная документация СУБД SQLite [Электронный ресурс]. URL: https://www.sqlite.org (дата обращения: 29.05.2026)."
    ]
    for idx, src in enumerate(lit_sources, 1):
        p_lit = doc.add_paragraph()
        p_lit.paragraph_format.first_line_indent = Inches(0.49)
        p_lit.paragraph_format.space_after = Pt(4)
        run_lit = p_lit.add_run(f"{idx}. {src}")
        run_lit.font.name = 'Times New Roman'
        run_lit.font.size = Pt(14)

    # ================= 10. ПРИЛОЖЕНИЯ =================
    add_chapter("10. Приложения")
    add_section("Приложение А. Основной файл логики приложения app.py")
    doc.add_paragraph("Ниже представлен листинг начала файла логики app.py:")
    
    app_py_start = (
        "from flask import Flask, render_template, request, redirect, url_for, flash, Response\n"
        "import sqlite3\n"
        "import csv\n"
        "import io\n\n"
        "app = Flask(__name__)\n"
        "app.secret_key = 'super_secret_key'\n\n"
        "def get_db_connection():\n"
        "    conn = sqlite3.connect('database.db')\n"
        "    conn.row_factory = sqlite3.Row\n"
        "    return conn"
    )
    
    p_app_code = doc.add_paragraph()
    p_app_code.paragraph_format.left_indent = Inches(0.25)
    p_app_code.paragraph_format.first_line_indent = Inches(0)
    p_app_code.paragraph_format.line_spacing = 1.0
    run_app_code = p_app_code.add_run(app_py_start)
    run_app_code.font.name = 'Courier New'
    run_app_code.font.size = Pt(10)

    # Сохранение документа
    doc.save("Курсовая_работа_Сборочный_цех.docx")
    print("Файл Курсовая_работа_Сборочный_цех.docx успешно обновлен по шаблону!")

if __name__ == "__main__":
    build_docx()
