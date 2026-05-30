import docx
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

TEMPLATES_DIR = os.path.join(os.path.dirname(__file__), 'templates_doc')
OUTPUTS_DIR = os.path.join(os.path.dirname(__file__), 'generated_docs')

def ensure_outputs_dir():
    os.makedirs(OUTPUTS_DIR, exist_ok=True)
    return OUTPUTS_DIR

def _set_cell_shading(cell, color_hex):
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear',
    })
    shading.append(shading_elm)

def _add_styled_paragraph(doc, text, bold=False, size=11, alignment=None, color=None, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if alignment is not None:
        p.alignment = alignment
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def _add_header_row(table, headers, color_hex='1e293b'):
    row = table.rows[0]
    for i, text in enumerate(headers):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.name = 'Calibri'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_cell_shading(cell, color_hex)

def _add_data_row(table, values):
    row = table.add_row()
    for i, val in enumerate(values):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(str(val))
        run.font.size = Pt(9)
        run.font.name = 'Calibri'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT

def _add_total_row(table, labels, color_hex='f1f5f9'):
    row = table.add_row()
    for i, text in enumerate(labels):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = 'Calibri'
        _set_cell_shading(cell, color_hex)

def fill_template(template_name, replacements, output_name):
    template_path = os.path.join(TEMPLATES_DIR, template_name)
    if not os.path.exists(template_path):
        raise FileNotFoundError(f'Шаблон не найден: {template_path}')
    
    doc = Document(template_path)
    
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            for key, value in replacements.items():
                placeholder = '{{ ' + key + ' }}'
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, str(value))
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        for key, value in replacements.items():
                            placeholder = '{{ ' + key + ' }}'
                            if placeholder in run.text:
                                run.text = run.text.replace(placeholder, str(value))
    
    out_path = os.path.join(ensure_outputs_dir(), output_name)
    doc.save(out_path)
    return out_path

def default_replacements():
    now = datetime.now()
    return {
        'date': now.strftime('%d.%m.%Y'),
        'time': now.strftime('%H:%M'),
        'datetime': now.strftime('%d.%m.%Y %H:%M'),
        'year': str(now.year),
    }

def build_full_report(parts_stock, products_stock, compositions, deficit_data, production_plan, output_name='full_report.docx'):
    doc = Document()
    
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2)
    section.right_margin = Cm(1.5)
    
    STYLE_HEADER = {'bold': True, 'size': 16, 'alignment': WD_ALIGN_PARAGRAPH.CENTER, 'color': (30, 41, 59), 'space_after': 2}
    STYLE_SUBHEADER = {'bold': False, 'size': 10, 'alignment': WD_ALIGN_PARAGRAPH.CENTER, 'color': (100, 116, 139), 'space_after': 12}
    STYLE_SECTION = {'bold': True, 'size': 12, 'color': (30, 41, 59), 'space_before': 12, 'space_after': 4}
    
    # Титул
    _add_styled_paragraph(doc, 'ООО «ТехноПромСборка»', **STYLE_HEADER)
    _add_styled_paragraph(doc, 'Сборочный цех — Система учета производства', **STYLE_SUBHEADER)
    _add_styled_paragraph(doc, f'КОМПЛЕКСНЫЙ ОТЧЁТ СОСТОЯНИЯ ПРОИЗВОДСТВА', bold=True, size=14,
                          alignment=WD_ALIGN_PARAGRAPH.CENTER, color=(79, 115, 150), space_before=6, space_after=2)
    now = datetime.now()
    _add_styled_paragraph(doc, f'Дата формирования: {now.strftime("%d.%m.%Y %H:%M")}', 
                          size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER, color=(100, 116, 139), space_after=18)
    
    # ========= Раздел 1: Детали =========
    _add_styled_paragraph(doc, '1. СПРАВКА О НАЛИЧИИ ДЕТАЛЕЙ НА СКЛАДЕ', **STYLE_SECTION)
    
    if parts_stock:
        table = doc.add_table(rows=1, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Light Shading Accent 1'
        _add_header_row(table, ['Ячейка', 'Деталь', 'Артикул', 'Количество (шт.)'])
        total_qty = 0
        for row in parts_stock:
            _add_data_row(table, [row['cell_number'], row['name'], row['article'], row['quantity_stored']])
            total_qty += row['quantity_stored']
        _add_total_row(table, [f'ИТОГО: {len(parts_stock)} ячеек', '', '', f'{total_qty} шт.'])
    else:
        _add_styled_paragraph(doc, 'Нет данных о деталях на складе.', size=10, color=(148, 163, 184))
    
    # ========= Раздел 2: Изделия =========
    _add_styled_paragraph(doc, '2. СПРАВКА О НАЛИЧИИ ГОТОВЫХ ИЗДЕЛИЙ', **STYLE_SECTION)
    
    if products_stock:
        table = doc.add_table(rows=1, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Light Shading Accent 1'
        _add_header_row(table, ['Ячейка', 'Изделие', 'Количество (шт.)'])
        total_qty = 0
        for row in products_stock:
            _add_data_row(table, [row['cell_number'], row['name'], row['quantity_stored']])
            total_qty += row['quantity_stored']
        _add_total_row(table, [f'ИТОГО: {len(products_stock)} ячеек', '', f'{total_qty} шт.'])
    else:
        _add_styled_paragraph(doc, 'Нет данных о готовых изделиях на складе.', size=10, color=(148, 163, 184))
    
    # ========= Раздел 3: Спецификации =========
    _add_styled_paragraph(doc, '3. СПРАВКА О КОМПЛЕКТЕ ДЕТАЛЕЙ (СПЕЦИФИКАЦИИ)', **STYLE_SECTION)
    
    if compositions:
        table = doc.add_table(rows=1, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Light Shading Accent 1'
        _add_header_row(table, ['Изделие', 'Деталь', 'Артикул', 'Кол-во в 1 шт.'])
        for row in compositions:
            _add_data_row(table, [row['product_name'], row['part_name'], row['article'], row['quantity']])
        _add_total_row(table, [f'Всего записей: {len(compositions)}', '', '', ''])
    else:
        _add_styled_paragraph(doc, 'Спецификации не заданы.', size=10, color=(148, 163, 184))
    
    # ========= Раздел 4: Дефицит =========
    _add_styled_paragraph(doc, '4. РАСЧЁТ ДЕФИЦИТА ДЕТАЛЕЙ', **STYLE_SECTION)
    
    _add_styled_paragraph(doc, 'Программа выпуска на месяц:', bold=True, size=10, space_before=4, space_after=4)
    for prod_name, plan_qty in production_plan:
        _add_styled_paragraph(doc, f'  • {prod_name}: {plan_qty} шт.', size=9, color=(71, 85, 105), space_after=1)
    
    doc.add_paragraph()  # blank line
    
    if deficit_data:
        table = doc.add_table(rows=1, cols=5)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Light Shading Accent 1'
        _add_header_row(table, ['Деталь', 'Артикул', 'Требуется', 'В наличии', 'Дефицит'])
        total_deficit = 0
        for row in deficit_data:
            deficit_val = row['deficit']
            is_deficit = deficit_val > 0
            _add_data_row(table, [row['part_name'], row['article'], row['total_needed'], row['in_stock'], deficit_val])
            total_deficit += deficit_val
            if is_deficit:
                for cell in table.rows[-1].cells:
                    _set_cell_shading(cell, 'fef2f2')
        _add_total_row(table, ['Общий дефицит:', '', '', '', f'{total_deficit} шт.'])
    else:
        _add_styled_paragraph(doc, 'Дефицит отсутствует.', size=10, color=(148, 163, 184))
    
    # Подвал
    doc.add_paragraph()
    _add_styled_paragraph(doc, '—' * 60, size=8, alignment=WD_ALIGN_PARAGRAPH.CENTER, color=(200, 200, 200))
    _add_styled_paragraph(doc, f'Отчёт сформирован: {now.strftime("%d.%m.%Y %H:%M")} | ООО «ТехноПромСборка»',
                          size=8, alignment=WD_ALIGN_PARAGRAPH.CENTER, color=(148, 163, 184))
    
    out_path = os.path.join(ensure_outputs_dir(), output_name)
    doc.save(out_path)
    return out_path


def build_purchase_order(missing_data, production_plan, output_name='purchase_order.docx'):
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(1.5)

    now = datetime.now()

    _add_styled_paragraph(doc, 'ООО «ТехноПромСборка»', bold=True, size=14,
                          alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    _add_styled_paragraph(doc, 'г. Рыбинск', size=9,
                          alignment=WD_ALIGN_PARAGRAPH.CENTER, color=(100, 116, 139), space_after=8)
    _add_styled_paragraph(doc, 'ЗАКАЗ НА ПОСТАВКУ ДЕТАЛЕЙ № ЗП-{0:04d}'.format(now.day * 100 + now.hour),
                          bold=True, size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                          color=(79, 115, 150), space_after=4)
    _add_styled_paragraph(doc, f'от «{now.day}» {now.strftime("%B %Y").lower()} г.',
                          size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                          color=(71, 85, 105), space_after=14)

    _add_styled_paragraph(doc, 'Поставщик: ___________________________________', size=10, space_after=2)
    _add_styled_paragraph(doc, 'Основание: Месячная программа выпуска изделий', size=10, space_after=12)

    _add_styled_paragraph(doc, 'Программа выпуска:', bold=True, size=10, space_after=4)
    for prod_name, plan_qty in production_plan:
        if plan_qty > 0:
            _add_styled_paragraph(doc, f'  • {prod_name}: {plan_qty} шт.', size=9, color=(71, 85, 105), space_after=1)

    doc.add_paragraph()

    if missing_data:
        table = doc.add_table(rows=1, cols=5)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Light Shading Accent 1'
        _add_header_row(table, ['№', 'Наименование детали', 'Артикул', 'Количество', 'Цена'])
        total_order = 0
        for i, row in enumerate(missing_data, 1):
            _add_data_row(table, [str(i), row['part_name'], row['article'], f'{row["to_order"]} шт.', '_____'])
            total_order += row['to_order']
            for cell in table.rows[-1].cells:
                _set_cell_shading(cell, 'fef2f2')
        _add_total_row(table, ['', 'ИТОГО к заказу:', '', f'{total_order} шт.', ''])

    doc.add_paragraph()
    doc.add_paragraph()

    _add_styled_paragraph(doc, 'Заказчик:', bold=True, size=10, space_before=8, space_after=20)
    _add_styled_paragraph(doc, 'Генеральный директор', size=9, color=(100, 116, 139), space_after=2)
    _add_styled_paragraph(doc, '__________________ /_______________/', size=10, space_after=2)
    _add_styled_paragraph(doc, 'М.П.', size=9, color=(148, 163, 184), space_after=14)

    _add_styled_paragraph(doc, '—' * 60, size=8, alignment=WD_ALIGN_PARAGRAPH.CENTER, color=(200, 200, 200))
    _add_styled_paragraph(doc, f'Сформировано: {now.strftime("%d.%m.%Y %H:%M")} | ООО «ТехноПромСборка»',
                          size=8, alignment=WD_ALIGN_PARAGRAPH.CENTER, color=(148, 163, 184))

    out_path = os.path.join(ensure_outputs_dir(), output_name)
    doc.save(out_path)
    return out_path
